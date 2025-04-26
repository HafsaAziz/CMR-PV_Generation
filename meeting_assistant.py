import os
import io
import base64
import streamlit as st
from PIL import Image
import tempfile
from pydub import AudioSegment
import subprocess
import time
import random
from google import generativeai as genai
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import requests
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
import json
from bs4 import BeautifulSoup

# Configuration des APIs
def configure_apis():
    # Configure Google Gemini
    google_api_key = st.secrets.get("GOOGLE_API_KEY")
    if google_api_key:
        genai.configure(api_key=google_api_key)
    else:
        st.error("❌ Clé API Google non trouvée!")
        st.stop()

def convert_to_mp3(input_path, output_path):
    """Convertit n'importe quel format audio en MP3"""
    try:
        import shutil
        if not shutil.which("ffmpeg"):
            st.error("ffmpeg non trouvé. Veuillez installer ffmpeg.")
            return False
        audio = AudioSegment.from_file(input_path)
        audio.export(output_path, format="mp3")
        return True
    except Exception as e:
        st.error(f"Erreur de conversion audio : {e}")
        return False

def extract_file_id_from_url(url):
    """Extrait l'ID du fichier depuis une URL Google Drive"""
    patterns = [
        r"https://drive\.google\.com/file/d/([a-zA-Z0-9_-]+)",
        r"https://drive\.google\.com/open\?id=([a-zA-Z0-9_-]+)",
        r"https://drive\.google\.com/uc\?id=([a-zA-Z0-9_-]+)",
        r"id=([a-zA-Z0-9_-]+)"
    ]
    
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            return match.group(1)
    return None

def verify_video_file(file_path):
    """Vérifie si le fichier vidéo est valide"""
    try:
        if not os.path.exists(file_path):
            st.error(f"❌ Le fichier {file_path} n'existe pas")
            return False
            
        file_size = os.path.getsize(file_path)
        if file_size < 10000:  # Moins de 10KB est suspect
            st.error("❌ Le fichier est trop petit pour être une vidéo valide")
            return False
            
        st.info(f"📊 Taille du fichier : {file_size/1024/1024:.1f} MB")
        
        # Vérification du format avec ffprobe
        probe_command = [
            "ffprobe",
            "-v", "error",
            "-show_format",
            "-show_streams",
            file_path
        ]
        
        result = subprocess.run(probe_command, capture_output=True, text=True)
        
        if result.returncode != 0:
            st.error(f"❌ Format vidéo non valide: {result.stderr}")
            return False
            
        st.success("✅ Format vidéo validé")
        return True
            
    except Exception as e:
        st.error(f"❌ Erreur lors de la vérification: {str(e)}")
        return False

def convert_vro_to_mp4(input_path, output_path):
    """Convertit un fichier VRO en MP4"""
    try:
        st.info("🔄 Conversion du fichier VRO en MP4...")
        
        # Commande de conversion optimisée pour les fichiers VRO
        convert_command = [
            "ffmpeg",
            "-y",  # Écraser le fichier de sortie si existant
            "-fflags", "+genpts",  # Générer les timestamps
            "-i", input_path,
            "-c:v", "libx264",  # Codec vidéo
            "-preset", "ultrafast",  # Conversion rapide
            "-crf", "23",  # Qualité raisonnable
            "-c:a", "aac",  # Codec audio
            "-strict", "experimental",
            "-b:a", "192k",  # Bitrate audio
            "-movflags", "+faststart",  # Optimisation pour la lecture web
            output_path
        ]
        
        # Exécuter la conversion
        result = subprocess.run(convert_command, capture_output=True, text=True)
        
        if result.returncode != 0:
            st.error(f"❌ Erreur lors de la conversion VRO: {result.stderr}")
            return False
            
        if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
            st.success("✅ Conversion VRO → MP4 réussie")
            return True
        else:
            st.error("❌ Fichier MP4 non créé ou vide")
            return False
            
    except Exception as e:
        st.error(f"❌ Erreur lors de la conversion: {str(e)}")
        return False

def extract_audio_from_video(input_video_path, output_audio_path):
    """Extrait l'audio d'une vidéo"""
    try:
        # Vérifier si le fichier existe
        if not os.path.exists(input_video_path):
            st.error("❌ Le fichier vidéo n'existe pas")
            return False
            
        # Vérifier si le fichier est vide
        if os.path.getsize(input_video_path) == 0:
            st.error("❌ Le fichier vidéo est vide")
            return False
            
        # Si c'est un fichier VRO, on le convertit d'abord en MP4
        if input_video_path.lower().endswith('.vro'):
            st.info("🔄 Conversion du fichier VRO en MP4...")
            temp_mp4 = input_video_path + '.mp4'
            try:
                # Commande de conversion VRO vers MP4
                convert_command = [
                    'ffmpeg',
                    '-i', input_video_path,
                    '-c:v', 'libx264',
                    '-preset', 'ultrafast',
                    '-c:a', 'aac',
                    '-strict', 'experimental',
                    '-write_xing', '0',
                    '-y',
                    temp_mp4
                ]
                
                result = subprocess.run(convert_command, capture_output=True, text=True)
                if result.returncode != 0:
                    st.error(f"❌ Erreur lors de la conversion VRO vers MP4: {result.stderr}")
                    return False
                    
                input_video_path = temp_mp4
                st.success("✅ Conversion VRO vers MP4 réussie")
                
            except Exception as e:
                st.error(f"❌ Erreur lors de la conversion VRO vers MP4: {str(e)}")
                return False
                
        # Extraction de l'audio
        st.info("🎵 Extraction de l'audio...")
        try:
            # Commande d'extraction audio
            extract_command = [
                'ffmpeg',
                '-i', input_video_path,
                '-vn',
                '-acodec', 'libmp3lame',
                '-ar', '44100',
                '-ab', '192k',
                '-y',
                output_audio_path
            ]
            
            result = subprocess.run(extract_command, capture_output=True, text=True)
            if result.returncode != 0:
                st.error(f"❌ Erreur lors de l'extraction audio: {result.stderr}")
                return False
                
            # Vérifier si le fichier audio a été créé et n'est pas vide
            if not os.path.exists(output_audio_path) or os.path.getsize(output_audio_path) == 0:
                st.error("❌ Le fichier audio n'a pas été créé ou est vide")
                return False
                
            st.success("✅ Extraction audio réussie")
            return True
            
        except Exception as e:
            st.error(f"❌ Erreur lors de l'extraction audio: {str(e)}")
            return False
            
    except Exception as e:
        st.error(f"❌ Erreur inattendue lors du traitement: {str(e)}")
        return False
    finally:
        # Nettoyage des fichiers temporaires
        try:
            temp_mp4 = input_video_path + '.mp4'
            if os.path.exists(temp_mp4):
                os.remove(temp_mp4)
        except Exception as e:
            st.warning(f"⚠️ Erreur lors du nettoyage des fichiers temporaires: {str(e)}")

def segment_audio(audio_path, segment_length_ms=120000):
    """Divise un gros fichier audio en segments sans tout charger en RAM"""
    try:
        import math
        import subprocess
        
        # Utiliser ffmpeg pour obtenir la durée totale
        result = subprocess.run([
            'ffprobe', '-v', 'error', '-show_entries',
            'format=duration', '-of',
            'default=noprint_wrappers=1:nokey=1', audio_path
        ], stdout=subprocess.PIPE, stderr=subprocess.STDOUT)
        
        total_duration = float(result.stdout)
        segment_length_sec = segment_length_ms / 1000
        
        num_segments = math.ceil(total_duration / segment_length_sec)
        segments = []
        
        for i in range(num_segments):
            start_time = i * segment_length_sec
            temp_segment_path = os.path.join(tempfile.gettempdir(), f"segment_{i+1}.mp3")
            
            # Extraire un petit segment avec ffmpeg sans tout charger
            extract_cmd = [
                "ffmpeg",
                "-y",
                "-i", audio_path,
                "-ss", str(start_time),
                "-t", str(segment_length_sec),
                "-c", "copy",
                temp_segment_path
            ]
            subprocess.run(extract_cmd, stdout=subprocess.DEVNULL, stderr=subprocess.STDOUT)
            
            if os.path.exists(temp_segment_path):
                segment_audio = AudioSegment.from_file(temp_segment_path)
                segments.append(segment_audio)
        
        return segments
        
    except Exception as e:
        st.error(f"❌ Erreur lors de la segmentation audio (stream) : {str(e)}")
        return []


def process_segment_batch(segments, start_idx, batch_size, total_segments, temp_dir, progress_bar, status_text):
    """Traite un lot de segments audio"""
    batch_transcript = []
    
    for i in range(start_idx, min(start_idx + batch_size, total_segments)):
        segment = segments[i]
        segment_number = i + 1
        
        try:
            segment_path = os.path.join(temp_dir, f"segment_{segment_number}.mp3")
            status_text.text(f"🎯 Traitement du segment {segment_number}/{total_segments}")
            
            segment.export(segment_path, format="mp3")
            
            with open(segment_path, "rb") as f:
                audio_bytes = f.read()

            model = genai.GenerativeModel('gemini-2.0-flash')
            response = model.generate_content([
                "Transcrivez ce segment audio mot pour mot en français.",
                {"mime_type": "audio/mp3", "data": audio_bytes}
            ])
            
            if response.text:
                batch_transcript.append(response.text)
                progress_bar.progress((i + 1)/total_segments)
                
        except Exception as e:
            st.warning(f"⚠️ Erreur sur le segment {segment_number}: {str(e)}")
            batch_transcript.append(f"[Segment {segment_number} non transcrit]")
            
        time.sleep(random.uniform(1, 2))
    
    return batch_transcript

def transcribe_video(video_file):
    """Transcrit une vidéo en texte sans charger tout en RAM."""
    with tempfile.TemporaryDirectory() as temp_dir:
        try:
            # Créer un chemin temporaire avec bonne extension
            if hasattr(video_file, 'name'):
                ext = os.path.splitext(video_file.name)[1].lower()
                video_temp_path = os.path.join(temp_dir, f"uploaded_video{ext}")
                st.info(f"📝 Sauvegarde du fichier temporaire: {video_file.name}")
            else:
                ext = '.mp4'
                video_temp_path = os.path.join(temp_dir, "uploaded_video.mp4")
                st.info("📝 Sauvegarde d'un fichier vidéo sans nom")

            # Maintenant, on enregistre directement l'objet téléchargé
            with open(video_temp_path, 'wb') as out_file:
                for chunk in iter(lambda: video_file.read(1024 * 1024), b''):
                    out_file.write(chunk)

            st.success("✅ Vidéo sauvegardée sur disque temporaire")

            # On vérifie la taille
            video_size = os.path.getsize(video_temp_path)
            st.info(f"📊 Taille du fichier vidéo: {video_size/1024/1024:.2f} MB")

            # Vérifier la validité
            if not verify_video_file(video_temp_path):
                return ""

            # On extrait maintenant l'audio
            audio_path = os.path.join(temp_dir, "output_audio.mp3")
            st.info("🎵 Extraction de l'audio...")
            if not extract_audio_from_video(video_temp_path, audio_path):
                return ""

            # Segmentation
            st.info("🔄 Segmentation de l'audio...")
            segments = segment_audio(audio_path)
            if not segments:
                st.error("❌ Échec de la segmentation audio")
                return ""

            st.success(f"✅ Audio segmenté en {len(segments)} parties")
            
            # Traitement par batch
            progress_bar = st.progress(0)
            status_text = st.empty()

            full_transcript = []
            BATCH_SIZE = 10

            for batch_start in range(0, len(segments), BATCH_SIZE):
                batch_results = process_segment_batch(
                    segments, batch_start, BATCH_SIZE, len(segments),
                    temp_dir, progress_bar, status_text
                )
                full_transcript.extend(batch_results)

            if not full_transcript:
                st.warning("⚠️ Aucun texte n'a été transcrit")
                return ""
                
            st.success("✅ Transcription terminée avec succès")
            return "\n".join(full_transcript)

        except Exception as e:
            st.error(f"❌ Erreur lors de la transcription: {str(e)}")
            return ""


def process_handwritten_image(image_bytes):
    """Extrait le texte d'une image manuscrite avec mécanisme de retry"""
    @retry_with_backoff
    def transcribe_image():
        try:
            image_base64 = base64.b64encode(image_bytes).decode('utf-8')
            
            model = genai.GenerativeModel('gemini-2.0-flash')
            
            prompt = """Transcris précisément le texte manuscrit dans cette image.
            INSTRUCTIONS :
            1. Retourne uniquement le texte, sans commentaires
            2. Préserve la mise en forme (retours à la ligne, espacements)
            3. Conserve la ponctuation exacte
            4. Maintiens les nombres et symboles tels quels
            5. Respecte les majuscules et minuscules"""
            
            response = model.generate_content([
                prompt,
                {"mime_type": "image/jpeg", "data": image_base64}
            ])
            
            if response.text:
                return response.text.strip()
            else:
                raise Exception("Aucun texte détecté dans l'image.")
                
        except Exception as e:
            st.warning(f"⚠️ Tentative de transcription échouée : {str(e)}")
            raise e

    try:
        # Premier essai
        result = transcribe_image()
        if result:
            return result
            
        # Si le résultat est vide, on attend et on réessaie
        time.sleep(2)  # Attente de 2 secondes
        st.info("🔄 Nouvelle tentative de transcription...")
        
        # Deuxième essai avec un prompt plus détaillé
        prompt_retry = """Analyse et transcris TOUT le texte manuscrit visible dans cette image.
        IMPORTANT :
        - Examine l'image en détail, pixel par pixel
        - Transcris absolument tout le texte visible
        - N'oublie aucun détail, même les petites annotations
        - Conserve la structure exacte du texte
        - Inclus les numéros, symboles et caractères spéciaux"""
        
        model = genai.GenerativeModel('gemini-2.0-flash')
        image_base64 = base64.b64encode(image_bytes).decode('utf-8')
        
        response = model.generate_content([
            prompt_retry,
            {"mime_type": "image/jpeg", "data": image_base64}
        ])
        
        if response.text:
            return response.text.strip()
        else:
            st.warning("⚠️ Aucun texte n'a été détecté dans l'image après plusieurs tentatives.")
            return ""
            
    except Exception as e:
        st.error(f"❌ Erreur lors de la reconnaissance du texte : {str(e)}")
        return ""

def retry_with_backoff(func, max_retries=5, initial_delay=1):
    """Fonction utilitaire pour réessayer une opération avec un délai exponentiel"""
    def wrapper(*args, **kwargs):
        delay = initial_delay
        last_exception = None
        
        for attempt in range(max_retries):
            try:
                return func(*args, **kwargs)
            except Exception as e:
                last_exception = e
                error_code = str(e)
                # Réessayer sur les erreurs de quota (429) ou d'annulation client (499)
                if "429" in error_code or "499" in error_code: 
                    st.warning(f"⚠️ Erreur API ({error_code}), nouvelle tentative {attempt + 1}/{max_retries} dans {delay} secondes...")
                    time.sleep(delay)
                    delay *= 2  # Backoff exponentiel
                else:
                    # Pour les autres exceptions, ne pas réessayer
                    raise e
        
        st.error(f"❌ Échec après {max_retries} tentatives : {str(last_exception)}")
        # Retourner None ou une valeur indiquant l'échec si toutes les tentatives échouent
        return None
    
    return wrapper

def process_pdf(pdf_file):
    """Extrait le contenu détaillé et les acronymes d'un PDF en un seul appel."""
    try:
        pdf_bytes = pdf_file.read()
        pdf_base64 = base64.b64encode(pdf_bytes).decode('utf-8')
        
        model = genai.GenerativeModel('gemini-2.0-flash')
        
        prompt = """Analyse ce document PDF de manière EXHAUSTIVE et DÉTAILLÉE.
        
        INSTRUCTIONS SPÉCIFIQUES :
        
        1. EXTRACTION COMPLÈTE DU CONTENU :
           - Extraire TOUS les textes, exactement comme ils apparaissent.
           - Conserver TOUS les chiffres, statistiques, données numériques avec leurs unités.
           - Maintenir TOUS les tableaux avec leurs données complètes.
           - Décrire TOUS les graphiques avec leurs valeurs précises.
           - Capturer TOUTES les notes de bas de page et références.
           - Respecter la structure (sections, titres, listes).
           - NE PAS résumer ou synthétiser le corps du texte.
           
        2. EXTRACTION DES ACRONYMES :
           - Identifier TOUS les acronymes présents dans le document.
           - Fournir leur définition complète telle qu'elle apparaît dans le texte (ou si elle est évidente).
           - Lister les acronymes et leurs définitions SÉPARÉMENT à la fin.
        
        3. FORMAT DE SORTIE ATTENDU :
           - D'abord, le contenu complet et détaillé du document, en respectant sa structure.
           - Ensuite, une ligne de séparation claire comme : '--- ACRONYMES ---'.
           - Enfin, la liste des acronymes, un par ligne, au format : 'ACRONYME: Définition complète'.
           
        IMPORTANT : Assure-toi de bien séparer le contenu principal de la liste des acronymes avec '--- ACRONYMES ---'."""
        
        @retry_with_backoff
        def analyze_pdf_and_extract_acronyms():
            response = model.generate_content([
                {
                    "role": "user",
                    "parts": [
                        prompt,
                        {"mime_type": "application/pdf", "data": pdf_base64}
                    ]
                }
            ])
            return response.text if response.text else ""
        
        full_result = analyze_pdf_and_extract_acronyms()
        
        if not full_result:
            st.warning(f"⚠️ Aucun contenu extrait du PDF: {pdf_file.name}")
            return {"summary": "", "acronyms": {}}
            
        # Séparer le contenu et les acronymes
        separator = "--- ACRONYMES ---"
        if separator in full_result:
            summary_part, acronym_part = full_result.split(separator, 1)
            summary = summary_part.strip()
            
            # Parser les acronymes
            acronyms = {}
            lines = acronym_part.strip().split('\n')
            for line in lines:
                if ':' in line:
                    acronym, definition = line.split(':', 1)
                    acronym = acronym.strip().upper()
                    definition = definition.strip()
                    if acronym and definition:
                        acronyms[acronym] = definition
            return {"summary": summary, "acronyms": acronyms}
        else:
            # Si le séparateur n'est pas trouvé, retourner tout comme résumé et pas d'acronymes
            st.warning(f"⚠️ Séparateur d'acronymes non trouvé dans l'analyse de {pdf_file.name}")
            return {"summary": full_result.strip(), "acronyms": {}}
            
    except Exception as e:
        st.error(f"❌ Erreur lors de l'analyse du PDF {pdf_file.name}: {str(e)}")
        return {"summary": f"[Erreur lors de l'analyse du PDF: {str(e)}]", "acronyms": {}}

def create_word_pv(content, logo_path=None):
    """Crée un document Word à partir du contenu du PV avec un formatage professionnel.
       Utilise les données PDF pré-analysées depuis st.session_state.pdf_data.
    """
    doc = Document()
    
    # Récupérer les données PDF et meeting_info depuis st.session_state
    pdf_data = st.session_state.get('pdf_data', {})
    meeting_info = st.session_state.get('meeting_info', {})

    # --- Début de la modification : Extraire et supprimer la section RECOMMANDATIONS --- 
    extracted_reco_text = "" # Initialiser la variable pour le texte des recos
    reco_marker = "--- RECOMMANDATIONS ---"
    if reco_marker in content:
        # Séparer le contenu principal de la section des recommandations
        main_content_part, reco_section = content.split(reco_marker, 1)
        content = main_content_part.strip() # Mettre à jour le contenu principal
        extracted_reco_text = reco_section.strip() # Stocker la section des recos
    # --- Fin de la modification ---
    
    # Style du document
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    # Fonction utilitaire pour formater les cellules de tableau
    def format_table_cell(cell, text, alignment=WD_ALIGN_PARAGRAPH.LEFT, bold=False):
        # Vide le contenu précédent de la cellule
        while len(cell.paragraphs) > 1:
            p = cell.paragraphs[-1]
            cell._element.remove(p._element)
        
        if not cell.paragraphs:
            p = cell.add_paragraph()
        else:
            p = cell.paragraphs[0]
        
        # Vide le contenu du paragraphe
        for run in p.runs:
            run.clear()
        
        # Ajoute le nouveau texte et applique le formatage
        run = p.add_run(text)
        run.bold = bold
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        p.alignment = alignment
    
    # En-tête avec logo et texte
    header_table = doc.add_table(rows=1, cols=3)
    header_table.style = 'Table Grid'
    
    # Colonne gauche (texte français)
    left_cell = header_table.cell(0, 0)
    left_text = left_cell.add_paragraph()
    left_text.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left_text.add_run("Royaume du Maroc\nCaisse Marocaine des\nRetraites\nConseil d'Administration\nComité d'Audit")
    
    # Colonne centrale (logo)
    center_cell = header_table.cell(0, 1)
    try:
        logo_run = center_cell.paragraphs[0].add_run()
        logo_run.add_picture(logo_path, width=Inches(1.5))
        center_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except:
        pass
    
    # Colonne droite (texte arabe)
    right_cell = header_table.cell(0, 2)
    right_text = right_cell.add_paragraph()
    right_text.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_text.add_run("المملكة المغربية\nالصندوق المغربي\nللتقاعد\nالمجلس الإداري\nلجنة التدقيق")
    
    # Ligne de séparation
    doc.add_paragraph().add_run("_" * 70)
    
    # Titre du PV
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Récupérer le numéro du PV depuis st.session_state
    pv_number = meeting_info.get('pv_number', '[Numéro PV]') # Valeur par défaut si non trouvé
    title_text = f"PROCÈS-VERBAL DÉTAILLÉ N° {pv_number} DE LA RÉUNION DU COMITÉ D'AUDIT"
    title_run = title.add_run(title_text)
    title_run.bold = True
    title_run.font.size = Pt(14)
    
    doc.add_paragraph()  # Espace
    
    # Tableau d'informations
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Table Grid'
    info_table.autofit = False
    
    # Récupérer les informations de la réunion depuis st.session_state
    info_rows = [
        ("Date", meeting_info.get('date', '')),
        ("Lieu", meeting_info.get('lieu', '')),
        ("Heure début", meeting_info.get('heure_debut', '')),
        ("Heure fin", meeting_info.get('heure_fin', ''))
    ]
    
    # Remplir le tableau d'informations
    for i, (label, value) in enumerate(info_rows):
        cells = info_table.rows[i].cells
        cells[0].text = label
        cells[1].text = value
        # Définir les largeurs
        cells[0].width = Inches(1.5)
        cells[1].width = Inches(4.5)
    
    doc.add_paragraph()  # Espace
    
    # Section "ÉTAIENT PRÉSENTS :"
    presents_title = doc.add_paragraph()
    presents_title.add_run("ÉTAIENT PRÉSENTS :").bold = True
    
    # Fonction pour créer un tableau de participants
    def create_participants_table(participants, section_title=None):
        if section_title:
            section_para = doc.add_paragraph()
            section_para.add_run(section_title).bold = True
        
        table = doc.add_table(rows=len(participants), cols=2)
        table.style = 'Table Grid'
        table.autofit = False
        
        for i, (name, title) in enumerate(participants):
            cells = table.rows[i].cells
            # Formater le nom avec un tiret
            cells[0].text = f"- {name}" if not name.startswith("-") else name
            cells[1].text = title
            # Définir les largeurs
            cells[0].width = Inches(3.0)
            cells[1].width = Inches(3.0)
        
        return table
    
    # Ajouter les participants par section
    participants_by_section = meeting_info.get('participants_by_section', {})
    
    # Ajouter les participants par section
    for section, participants in participants_by_section.items():
        if participants:
            doc.add_paragraph()  # Espace avant la section
            create_participants_table(participants, section)
    
    doc.add_paragraph()
    
    # Traitement du contenu principal
    sections = content.split('\n\n')
    ordre_du_jour_processed = False
    in_ordre_du_jour = False
    
    for section in sections:
        section_stripped = section.strip()
        if not section_stripped:
            continue
            
        # Traiter l'ordre du jour une seule fois
        if not ordre_du_jour_processed and "ORDRE DU JOUR" in section.upper():
            in_ordre_du_jour = True
            ordre_du_jour_processed = True
            
            # Titre "ORDRE DU JOUR :"
            p_title = doc.add_paragraph()
            run_title = p_title.add_run("ORDRE DU JOUR :")
            run_title.font.name = 'Times New Roman'
            run_title.font.size = Pt(12)
            run_title.bold = True
            p_title.paragraph_format.space_after = Pt(12)
            
            continue
            
        # Traiter les points de l'ordre du jour
        if in_ordre_du_jour:
            if re.match(r'^\d+\.', section_stripped):
                p_point = doc.add_paragraph()
                p_point.paragraph_format.left_indent = Inches(0.5)
                p_point.paragraph_format.space_before = Pt(0)
                p_point.paragraph_format.space_after = Pt(0)
                run_point = p_point.add_run(section_stripped)
                run_point.font.name = 'Times New Roman'
                run_point.font.size = Pt(12)
            elif "L'ordre du jour proposé" in section_stripped:
                in_ordre_du_jour = False
                # Ajouter un paragraphe vide avant
                doc.add_paragraph()
                
                # Phrase de transition
                p_transition = doc.add_paragraph()
                p_transition.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_transition.paragraph_format.left_indent = Inches(0)
                run_transition = p_transition.add_run(section_stripped)
                run_transition.font.name = 'Times New Roman'
                run_transition.font.size = Pt(12)
                
                # Ajouter un paragraphe vide après
                doc.add_paragraph()
            continue
            
        # Traitement des sections principales (numérotées)
        if re.match(r"^\d+\.\s", section_stripped):
            doc.add_paragraph()  # Espace avant nouvelle section
            p = doc.add_paragraph()
            run = p.add_run(section_stripped)
            run.bold = True
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            continue # Passer à la section suivante

        # Traiter les tableaux uniquement s'ils ne sont pas des recommandations
        if '|' in section and not any(marker in section.lower() for marker in ['recommandation', 'recommendation']):
            # Détecter et créer un tableau
            rows = [row.strip() for row in section.split('\n') if '|' in row and not row.strip().startswith('|-')]
            if rows:
                try:
                    num_cols = len(rows[0].split('|')) - 2
                    if num_cols <= 0:
                        print(f"[WARN] Table dynamique détectée avec {num_cols} colonnes. Ignorée.")
                        continue
                    
                    table = doc.add_table(rows=len(rows), cols=num_cols)
                    table.style = 'Table Grid'
                    table.autofit = False
                    
                    # Calculer les largeurs de colonnes
                    total_width = 6.0
                    col_width = total_width / num_cols
                    
                    # Appliquer les largeurs et remplir le tableau
                    for i, row_text in enumerate(rows):
                        try:
                            cells_content = [cell.strip() for cell in row_text.split('|')[1:-1]]
                            if len(cells_content) != num_cols:
                                print(f"[WARN] Ligne {i} table dynamique a {len(cells_content)} cellules, attendu {num_cols}. Ligne ignorée.")
                                continue

                            for j, cell_content in enumerate(cells_content):
                                current_cell = table.rows[i].cells[j]
                                if current_cell is None:
                                    print(f"[ERROR] Cellule ({i},{j}) est None dans table dynamique. Cellule ignorée.")
                                    continue
                                
                                format_table_cell(current_cell, cell_content,
                                               alignment=WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT,
                                               bold=i == 0)
                                current_cell.width = Inches(col_width)
                        except Exception as e:
                            print(f"[ERROR] Erreur lors du traitement de la ligne {i}: {str(e)}")
                            continue
                except Exception as e:
                    print(f"[ERROR] Erreur lors de la création du tableau: {str(e)}")
                    continue

                doc.add_paragraph()  # Espace après le tableau
                continue
        
        # Gérer le texte normal et les listes à puces
        lines = section.split('\n')
        for line in lines:
            line_text = line.strip()
            if not line_text:
                continue

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)

            # Détecter et formater les listes à puces
            if line_text.startswith(('-', '*', '•')):
                text_content = re.sub(r"^[-*•]\s*", "", line_text)
                p.text = text_content
                p.style = 'List Bullet'
            else:
                p.text = line_text
                p.paragraph_format.first_line_indent = Inches(0.3)
    
    # 1. Tableau des recommandations
    doc.add_paragraph() # Ajoute un espace avant le titre
    recommendations_title = doc.add_paragraph()
    run_reco_title = recommendations_title.add_run("RECOMMANDATIONS")
    run_reco_title.bold = True
    run_reco_title.font.name = 'Times New Roman'
    run_reco_title.font.size = Pt(12)
    recommendations_title.paragraph_format.space_before = Pt(6) # Espace avant le titre
    recommendations_title.paragraph_format.space_after = Pt(6) # Espace après le titre

    # Créer la structure du tableau des recommandations (juste l'en-tête initialement)
    recommendations_table = doc.add_table(rows=1, cols=4)
    recommendations_table.style = 'Table Grid'
    recommendations_table.autofit = False
    
    # Définir les en-têtes
    headers = ["Domaine", "Recommandations", "Structure\nresponsable", "Échéance"]
    header_cells = recommendations_table.rows[0].cells
    for i, header in enumerate(headers):
        format_table_cell(header_cells[i], header, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    
    # Définir les largeurs des colonnes (total: 6 pouces)
    widths = [1.5, 2.5, 1.0, 1.0]  # en pouces
    for i, width in enumerate(widths):
        for cell in recommendations_table.columns[i].cells:
            cell.width = Inches(width)
    
    # --- Extraire et ajouter les recommandations --- 
    recommendations_data = []
    # Utiliser directement extracted_reco_text au lieu de chercher à nouveau dans content
    if extracted_reco_text:
        # Regex plus tolérante et recherche globale
        reco_pattern = re.compile(
            r"\[RECO\]\s*Domaine\s*=\s*\"(.*?)\"\s*\|\s*Recommandation\s*=\s*\"(.*?)\"\s*\|\s*Responsable\s*=\s*\"(.*?)\"\s*\|\s*Échéance\s*=\s*\"(.*?)\"",
            re.IGNORECASE | re.DOTALL
        )
        matches = reco_pattern.findall(extracted_reco_text)
        for match in matches:
            recommendations_data.append({
                "Domaine": match[0].strip(),
                "Recommandations": match[1].strip(),
                "Responsable": match[2].strip(),
                "Échéance": match[3].strip()
            })
        # Log les lignes qui commencent par [RECO] mais ne matchent pas
        for line in extracted_reco_text.strip().split('\n'):
            if line.strip().startswith('[RECO]') and not reco_pattern.match(line.strip()):
                print(f"[WARN] Ligne de recommandation non reconnue: {line.strip()}")

    # Remplacer le contenu traité pour ne plus inclure la section reco
    # content = main_content_for_later # Supprimé, content est déjà propre

    # Ajouter les lignes au tableau
    if recommendations_data:
        for reco in recommendations_data:
            row_cells = recommendations_table.add_row().cells
            format_table_cell(row_cells[0], reco.get("Domaine", "N/A"))
            format_table_cell(row_cells[1], reco.get("Recommandations", "N/A"))
            format_table_cell(row_cells[2], reco.get("Responsable", "N/A"))
            format_table_cell(row_cells[3], reco.get("Échéance", "N/A"))
            # Réappliquer les largeurs aux nouvelles cellules
            for i, width in enumerate(widths):
                 row_cells[i].width = Inches(width)
    else:
        # Ajouter une ligne indiquant "Aucune recommandation"
        row_cells = recommendations_table.add_row().cells
        # Écrire le message dans la première cellule, laisser les autres vides
        format_table_cell(row_cells[0], "Aucune recommandation identifiée", alignment=WD_ALIGN_PARAGRAPH.CENTER)
        # Optionnel: laisser les autres cellules vides ou mettre "-"
        format_table_cell(row_cells[1], "")
        format_table_cell(row_cells[2], "")
        format_table_cell(row_cells[3], "")
        # Réappliquer les largeurs aux nouvelles cellules
        for i, width in enumerate(widths):
             row_cells[i].width = Inches(width)

    # Fin de la section Recommandations
    doc.add_paragraph() # Espace après le tableau

    # 2. Annexes et références (Utilise pdf_data)
    doc.add_paragraph()
    annexes_title = doc.add_paragraph()
    annexes_title.add_run("ANNEXES ET RÉFÉRENCES").bold = True
    
    pdf_filenames = list(pdf_data.keys())
    if pdf_filenames:
        annexes_table = doc.add_table(rows=len(pdf_filenames), cols=2)
        annexes_table.style = 'Table Grid'
        annexes_table.autofit = False
        
        for i, filename in enumerate(pdf_filenames):
            cells = annexes_table.rows[i].cells
            format_table_cell(cells[0], f"Document : {i+1}")
            format_table_cell(cells[1], filename)
            cells[0].width = Inches(1.5)
            cells[1].width = Inches(4.5)
    else:
        annexes_table = doc.add_table(rows=1, cols=2)
        annexes_table.style = 'Table Grid'
        annexes_table.autofit = False
        cells = annexes_table.rows[0].cells
        format_table_cell(cells[0], "Aucun document annexe")
        format_table_cell(cells[1], "")
    
    doc.add_paragraph()  # Espace après les annexes
    
    # 3. Lexique technique (Utilise pdf_data)
    doc.add_paragraph()
    lexique_title = doc.add_paragraph()
    lexique_title.add_run("LEXIQUE TECHNIQUE").bold = True
    
    # Agréger tous les acronymes de tous les PDFs
    all_acronyms = {}
    for data in pdf_data.values():
        if isinstance(data, dict) and 'acronyms' in data:
             all_acronyms.update(data['acronyms']) # update fusionne les dictionnaires
    
    # Créer le tableau du lexique
    if all_acronyms:
        # Trier les acronymes par ordre alphabétique
        sorted_acronyms = sorted(all_acronyms.items())
        
        lexique_table = doc.add_table(rows=len(sorted_acronyms), cols=2)
        lexique_table.style = 'Table Grid'
        lexique_table.autofit = False
        
        for i, (acronym, definition) in enumerate(sorted_acronyms):
            cells = lexique_table.rows[i].cells
            format_table_cell(cells[0], acronym, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            format_table_cell(cells[1], definition)
            cells[0].width = Inches(1.5)
            cells[1].width = Inches(4.5)
    else:
        lexique_table = doc.add_table(rows=1, cols=2)
        lexique_table.style = 'Table Grid'
        lexique_table.autofit = False
        cells = lexique_table.rows[0].cells
        format_table_cell(cells[0], "Aucun acronyme trouvé")
        format_table_cell(cells[1], "")

    # Sauvegarder dans un buffer
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    
    return doc_buffer

def generate_meeting_minutes(video_transcript, handwritten_text, pdf_summary, meeting_info):
    """Génère un PV de réunion structuré avec un niveau de détail élevé et des données précises"""
    try:
        # Formater les sources d'information de manière plus structurée
        combined_text = ""
        
        # 1. Ajouter la transcription vidéo si disponible
        if video_transcript and video_transcript.strip():
            combined_text += "[TRANSCRIPTION VIDÉO]\n"
            combined_text += video_transcript.strip() + "\n\n"
        
        # 2. Ajouter les notes manuscrites si disponibles
        if handwritten_text and handwritten_text.strip():
            combined_text += "[NOTES MANUSCRITES]\n"
            combined_text += handwritten_text.strip() + "\n\n"
        
        # 3. Ajouter le contenu des documents PDF si disponible
        if pdf_summary and pdf_summary.strip():
            combined_text += "[DOCUMENTS PDF]\n"
            combined_text += pdf_summary.strip() + "\n\n"

        if not combined_text.strip():
            return "Aucun contenu disponible pour générer le PV."

        model = genai.GenerativeModel('gemini-2.0-flash')
        
        prompt = f"""Génère un procès-verbal détaillé et professionnel en utilisant TOUTES les sources d'information fournies.
        
        ⚠️ RÈGLES CRUCIALES :
        1. ABSOLUMENT PAS D'HALLUCINATIONS :
           - Ne générer QUE du contenu présent dans les sources fournies
           - Ne PAS inventer de faits, chiffres ou discussions
           - Ne PAS extrapoler ou ajouter des informations non présentes
           - En cas de doute, omettre plutôt qu'inventer

        2. EXACTITUDE DES NOMS DES PARTICIPANTS :
           - Utiliser UNIQUEMENT les noms des participants listés au début du PV
           - Vérifier l'orthographe EXACTE de chaque nom mentionné
           - Utiliser systématiquement le même format pour les titres (M., Mme)
           - Ne JAMAIS mentionner de participants non listés initialement
           - En cas de doute sur l'identité d'un intervenant, utiliser "un participant" plutôt que de risquer une erreur
           - Liste des participants autorisés et leurs titres :
           {meeting_info.get('participants_text', '')}
        
        3. MISE EN FORME SIMPLE ET EFFICACE :
           - Pas d'indentations inutiles dans le texte
           - Aligner tout le texte à gauche sauf indication contraire
           - Utiliser les retours à la ligne uniquement quand nécessaire
           - Éviter toute mise en forme décorative sans utilité
        
        SOURCES D'INFORMATION À INTÉGRER :
        1. Transcription de la vidéo de la réunion
        2. Notes manuscrites prises pendant la réunion
        3. Documents PDF analysés
        
        INSTRUCTIONS SPÉCIFIQUES POUR L'ORDRE DU JOUR :
        1. Commencer par une section "ORDRE DU JOUR :" seule sur sa ligne
        2. Lister les points avec ce format EXACT :
           - Un point par ligne (pas de ligne vide entre les points)
           - Numérotation : "1. ", "2. ", etc. (avec un espace après le point)
           - Texte en minuscules (sauf première lettre et noms propres)
           - Pas de formatage spécial (pas de gras, pas d'italique)
        3. Après le dernier point, sauter une ligne
        4. Ajouter la phrase de transition sur un nouveau paragraphe :
           "L'ordre du jour proposé ayant été adopté à l'unanimité, les membres du Comité présents ont entamé l'examen des points inscrits."
        
        Exemple EXACT du format attendu :
        ORDRE DU JOUR :
        1. Validation du procès-verbal de la réunion précédente
        2. Discussion sur les placements immobiliers et les fonds de société
        3. Point sur l'avancement du projet X

        L'ordre du jour proposé ayant été adopté à l'unanimité, les membres du Comité présents ont entamé l'examen des points inscrits.

        INSTRUCTIONS POUR LE RESTE DU CONTENU :
        - Utiliser la transcription vidéo comme source principale pour les discussions et interventions
        - Intégrer les détails des notes manuscrites pour compléter ou clarifier les points discutés
        - Incorporer les données et statistiques des documents PDF de manière contextuelle
        - Assurer la cohérence entre les différentes sources d'information
        - En cas de divergence entre les sources, privilégier dans l'ordre :
          1) La transcription vidéo (source primaire des discussions)
          2) Les notes manuscrites (annotations et précisions en temps réel)
          3) Les documents PDF (informations de référence)

        RÈGLES DE FORMATAGE ET CONTENU :
        
        1. FORMAT DE L'ORDRE DU JOUR :
           - Commencer par "ORDRE DU JOUR :"
           - Liste immédiate des points numérotés sans espaces entre eux et en minuscule sauf la premiere lettre et pas en gras.
           - Format exact attendu :
           ORDRE DU JOUR :
           1. VALIDATION DU PROCÈS-VERBAL DE LA RÉUNION PRÉCÉDENTE
           2. EXAMEN DES COMPTES DU PREMIER SEMESTRE 2024
           3. DISCUSSION SUR LES PLACEMENTS IMMOBILIERS ET LES FONDS DE SOCIÉTÉ
           4. POINT SUR L'AVANCEMENT DU PROJET X
        
           Après les points de l'ordre du jour, ajouter EXACTEMENT cette phrase sur une nouvelle ligne apres un saut de ligne:
           "L'ordre du jour proposé ayant été adopté à l'unanimité, les membres du Comité présents ont entamé l'examen des points inscrits."
        
        2. RÈGLES STRICTES POUR L'ORDRE DU JOUR :
           - PAS d'introduction ou de texte avant l'ordre du jour
           - PAS d'espace entre "ORDRE DU JOUR :" et le premier point
           - PAS d'espace entre les points
           - Numérotation simple : "1. ", "2. ", etc.
           - Texte des points en minuscule sauf la premiere lettre
           - Points alignés sans indentation
           - APRÈS les points, ajouter la phrase de transition EXACTEMENT comme spécifiée
        
        3. PRÉSENTATION DES DONNÉES ET STATISTIQUES :
           - Intégrer naturellement les statistiques importantes dans les paragraphes
           - Mettre en évidence les chiffres clés dans le contexte
           - Inclure les comparaisons et évolutions pertinentes
           - NE PAS INVENTER de chiffres ou statistiques non présents dans les sources
        
        4. STRUCTURE DU CONTENU APRÈS L'ORDRE DU JOUR :
           - Le corps du PV DOIT être structuré en sections distinctes, correspondant EXACTEMENT à chaque point de l'ordre du jour
           - Chaque section DOIT commencer par le numéro et le titre exact du point de l'ordre du jour (en majuscules)
           - Sous chaque titre de section, développer UNIQUEMENT les discussions, décisions et informations présentes dans les sources
           - NE PAS mélanger les informations de différents points
           - Assurer une transition logique et claire entre les sections
           - Tout le texte aligné à gauche sans indentation inutile
        
        5. RÈGLES DE RÉDACTION POUR UN PV PRÉCIS :
           - Style professionnel et formel
           - Phrases complètes et précises
           - Se limiter STRICTEMENT aux informations présentes dans les sources
           - Utiliser des marqueurs de liste simples si nécessaire
           - Assurer la précision absolue dans la présentation des faits
           - Éviter toute spéculation ou interprétation personnelle
           
        6. IDENTIFICATION ET FORMATAGE DES RECOMMANDATIONS :
           - Si des recommandations sont mentionnées dans TOUTE source (vidéo, notes, documents), les identifier
           - LIMITER le nombre total de recommandations à UN MAXIMUM DE 5 (les plus importantes uniquement)
           - À la fin du texte, ajouter une section '--- RECOMMANDATIONS ---'
           - Format pour chaque recommandation : 
             [RECO] Domaine="[domaine]" | Recommandation="[texte]" | Responsable="CMR" | Échéance="[délai]"
           - Ne PAS créer de recommandations non explicitement mentionnées dans les sources
           - IMPORTANT pour le champ Responsable :
             * TOUJOURS utiliser "CMR" comme structure responsable
             * NE JAMAIS mettre le nom d'un employé ou d'une personne
             * NE JAMAIS utiliser de sous-divisions ou de services spécifiques
             * La responsabilité est TOUJOURS attribuée à l'institution CMR dans son ensemble

        TRAITEMENT DES ANNEXES ET RÉFÉRENCES :
        1. IDENTIFICATION DES ANNEXES :
           - Identifier tous les documents mentionnés dans les sources
           - Rechercher leur signification et description exacte dans les sources
           - Inclure le titre complet et la référence précise de chaque document
           - Pour chaque annexe citée, vérifier :
             * Son titre officiel complet
             * Sa référence ou numéro si mentionné
             * Sa description ou son contenu principal tel que décrit dans les sources
           - Ne pas inventer de descriptions si non trouvées dans les sources

        2. CITATION DES ANNEXES DANS LE TEXTE :
           - Lors de la première mention d'une annexe, inclure sa référence complète
           - Utiliser la formulation exacte trouvée dans les sources
           - Si un document est mentionné sans description claire, utiliser uniquement son titre sans interprétation

        IMPORTANT :
        - Commencer DIRECTEMENT par "ORDRE DU JOUR :"
        - Maintenir un format EXACT pour l'ordre du jour
        - Utiliser UNIQUEMENT des majuscules pour les points de l'ordre du jour
        - Structurer le PV selon les points de l'ordre du jour
        - NE JAMAIS inventer ou extrapoler des informations
        - Éviter toute indentation ou mise en forme inutile"""

        @retry_with_backoff
        def generate_content():
            response = model.generate_content([
                {
                    "role": "user",
                    "parts": [f"""Analyse TOUTES les sources d'information suivantes et génère un PV détaillé et professionnel.
                    Assure-toi d'intégrer les informations de CHAQUE source de manière cohérente.

Sources d'information :
{combined_text}

Instructions Détaillées :
{prompt}"""]
                }
            ])
            return response.text if response.text else ""

        result = generate_content()
        
        if result:
            # Nettoyage et formatage du texte généré
            result = result.replace('**', '')
            result = result.replace('*', '')
            
            # Formater les titres de section
            for i in range(1, 10):
                result = result.replace(f'{i}.\n', f'{i}. ')
                result = result.replace(f'\n{i}. \n', f'\n{i}. ')
                result = result.replace(f'\n{i}.\n', f'\n{i}. ')
                result = result.replace(f'\n{i}.', f'\n\n{i}.')
            
            # Formater l'ordre du jour
            if "ORDRE DU JOUR" in result:
                ordre_index = result.index("ORDRE DU JOUR")
                result = result[:ordre_index] + "\n\nORDRE DU JOUR :\n" + result[ordre_index + 13:]
            
            # Formater les listes à puces
            result = result.replace('• ', '\n• ')
            
            # Assurer des sauts de ligne appropriés
            result = result.replace('\n\n\n', '\n\n')
            
            # Vérification finale pour les numéros isolés
            result = re.sub(r'\n(\d+\.)\s*\n', r'\n\1 ', result)
            
            return result.strip()
        else:
            st.warning("⚠️ Aucun contenu n'a été généré pour le PV.")
            return ""
                
    except Exception as e:
        st.error(f"❌ Erreur lors de la génération du PV : {str(e)}")
        return ""

def download_video_from_drive(video_url, output_path):
    """Télécharge une vidéo depuis Google Drive avec gestion des gros fichiers"""
    try:
        st.info("🔄 Initialisation du téléchargement...")
        
        # Extraire l'ID du fichier
        file_id = extract_file_id_from_url(video_url)
        if not file_id:
            st.error("❌ Format d'URL Google Drive non reconnu")
            return False

        st.info(f"📝 ID du fichier extrait : {file_id}")

        # Configuration de la session avec des headers complets
        session = requests.Session()
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'fr,fr-FR;q=0.8,en-US;q=0.5,en;q=0.3',
            'Accept-Encoding': 'gzip, deflate, br',
            'Connection': 'keep-alive',
            'Upgrade-Insecure-Requests': '1'
        }

        # Utiliser l'URL de téléchargement direct avec usercontent
        download_url = f'https://drive.usercontent.google.com/download?id={file_id}&export=download&authuser=0&confirm=t'
        st.info(f"🔍 Tentative de téléchargement direct : {download_url}")
        
        response = session.get(download_url, headers=headers, stream=True, timeout=30)
        st.info(f"📡 Code de statut : {response.status_code}")
        st.info(f"📝 Type de contenu : {response.headers.get('Content-Type', 'Non spécifié')}")

        # Vérifier si nous avons reçu un fichier et non une page HTML
        content_type = response.headers.get('Content-Type', '').lower()
        if 'text/html' in content_type:
            st.warning("⚠️ Redirection vers la page de confirmation détectée")
            # Essayer l'URL alternative pour les gros fichiers
            download_url = f'https://drive.usercontent.google.com/download?id={file_id}&export=download&authuser=0&confirm=t&uuid=123&at=123'
            st.info(f"🔍 Tentative avec URL pour gros fichiers : {download_url}")
            response = session.get(download_url, headers=headers, stream=True, timeout=30)
            
            content_type = response.headers.get('Content-Type', '').lower()
            if 'text/html' in content_type:
                st.error("❌ Impossible d'accéder au fichier. Assurez-vous que :")
                st.error("1. Le fichier est partagé avec 'Tout le monde avec le lien'")
                st.error("2. Vous avez les droits 'Lecteur' sur le fichier")
                st.error("3. Le fichier n'est pas dans la corbeille")
                return False

        # Utiliser un nom de fichier temporaire unique
        import tempfile
        temp_dir = tempfile.gettempdir()
        temp_path = os.path.join(temp_dir, f"download_{file_id}_{int(time.time())}.tmp")
        
        try:
            # Télécharger avec barre de progression
            progress_bar = st.progress(0)
            chunk_size = 500 * 1024 * 1024  # 500MB chunks
            downloaded_size = 0
            
            # Obtenir la taille réelle du fichier depuis les headers
            expected_size = None
            if 'content-length' in response.headers:
                expected_size = int(response.headers['content-length'])
                st.info(f"📦 Taille totale du fichier : {expected_size/1024/1024:.1f} MB")
            
            with open(temp_path, 'wb') as f:
                for chunk in response.iter_content(chunk_size=chunk_size):
                    if chunk:
                        f.write(chunk)
                        downloaded_size += len(chunk)
                        # Afficher la progression
                        current_mb = downloaded_size/1024/1024
                        if expected_size:
                            current_percent = (downloaded_size/expected_size) * 100
                            st.info(f"📥 Téléchargé : {current_mb:.1f} MB ({current_percent:.1f}%)")
                            # Mise à jour de la barre de progression
                            progress_bar.progress(min(1.0, downloaded_size/expected_size))
                        else:
                            # Si on n'a pas la taille totale, afficher juste la taille téléchargée
                            st.info(f"📥 Téléchargé : {current_mb:.1f} MB")

            # Vérifier le fichier téléchargé
            if os.path.exists(temp_path):
                file_size = os.path.getsize(temp_path)
                if file_size < 10000:  # Moins de 10KB
                    st.error("❌ Fichier téléchargé invalide ou trop petit")
                    try:
                        os.remove(temp_path)
                    except:
                        pass
                    return False
                
                # Vérifier les premiers octets pour s'assurer que c'est un fichier VRO
                with open(temp_path, 'rb') as f:
                    header = f.read(8)
                    if not header.startswith(b'DVD') and not header.startswith(b'\x00\x00\x01\xBA'):
                        st.error("❌ Le fichier téléchargé n'est pas un fichier VRO valide")
                        try:
                            os.remove(temp_path)
                        except:
                            pass
                        return False
                
                # Renommer le fichier temporaire
                try:
                    if os.path.exists(output_path):
                        os.remove(output_path)
                    os.rename(temp_path, output_path)
                except Exception as e:
                    st.error(f"❌ Erreur lors du déplacement du fichier : {str(e)}")
                    # Essayer de copier le fichier à la place
                    import shutil
                    try:
                        shutil.copy2(temp_path, output_path)
                        os.remove(temp_path)
                    except Exception as e2:
                        st.error(f"❌ Échec de la copie du fichier : {str(e2)}")
                        return False
                
                st.success(f"✅ Téléchargement réussi - Taille : {file_size/1024/1024:.1f} MB")
                
                # Convertir si c'est un VRO
                if output_path.lower().endswith('.vro'):
                    st.info("🔄 Conversion du fichier VRO en MP4...")
                    mp4_path = output_path + '.mp4'
                    if convert_vro_to_mp4(output_path, mp4_path):
                        try:
                            os.remove(output_path)
                            os.rename(mp4_path, output_path)
                            st.success("✅ Conversion VRO → MP4 réussie")
                            return True
                        except Exception as e:
                            st.error(f"❌ Erreur lors du renommage après conversion : {str(e)}")
                            return False
                    else:
                        st.error("❌ Échec de la conversion VRO")
                        if os.path.exists(mp4_path):
                            try:
                                os.remove(mp4_path)
                            except:
                                pass
                        return False
                
                return True
            else:
                st.error("❌ Échec de l'écriture du fichier")
                return False

        except Exception as e:
            st.error(f"❌ Erreur pendant le téléchargement : {str(e)}")
            try:
                if os.path.exists(temp_path):
                    os.remove(temp_path)
            except:
                pass
            return False

    except Exception as e:
        st.error(f"❌ Erreur inattendue : {str(e)}")
        try:
            if os.path.exists(output_path):
                os.remove(output_path)
        except:
            pass
        return False

def main():
    st.set_page_config(
        page_title="Assistant de Réunion CMR",
        page_icon="📊",
        layout="wide"
    )

    # Configuration de l'API Gemini
    configure_apis()
    
    # Ajout du titre avec logo
    st.markdown("""
        <div style='display: flex; align-items: center; justify-content: center; margin-bottom: 2rem;'>
            <img src='logo.png' style='height: 80px; margin-right: 20px;'>
            <h1 style='margin: 0;'>Génération de PV</h1>
        </div>
    """, unsafe_allow_html=True)
    
    # Variables de session pour stocker les résultats
    if 'video_transcript' not in st.session_state:
        st.session_state.video_transcript = ""
    if 'handwritten_text' not in st.session_state:
        st.session_state.handwritten_text = ""
    if 'pdf_summary' not in st.session_state:
        st.session_state.pdf_summary = ""
    if 'pdf_data' not in st.session_state:
        st.session_state.pdf_data = {}
    if 'meeting_info' not in st.session_state:
        st.session_state.meeting_info = None
    if 'additional_participants' not in st.session_state:
        st.session_state.additional_participants = []

    # Section d'informations de base du PV
    st.header("📝 Informations de base du PV")
    
    col1, col2 = st.columns(2)
    
    with col1:
        pv_number = st.text_input("Numéro du PV", "02/24")
        date = st.date_input("Date", format="DD/MM/YYYY")
        lieu = st.text_input("Lieu", "Salle du Conseil CMR")
        heure_debut = st.time_input("Heure début")
        heure_fin = st.time_input("Heure fin")
    
    with col2:
        st.subheader("Participants")
        participants = []
        
        # Sections prédéfinies avec leurs participants
        sections = {
            "Membres du Comité d'Audit": [
                ("M. Aziz LOUBANI", "Président du Comité d'Audit, Représentant du Ministère de l'Économie et des Finances"),
                ("M. Mustapha KASSI", "Expert et membre indépendant"),
                ("M. Mohammed EL HAJJOUJI", "Expert et membre indépendant")
            ],
            "Caisse Marocaine des Retraites": [
                ("M. Lotfi BOUJENDAR", "Directeur de la CMR"),
                ("M. Mohamed El Mokhtar LOUTFI", "Secrétaire Général de la CMR"),
                ("M. Mohamed Jaber KHEMLICHI", "Chef de Pôle Gestion de Portefeuille"),
                ("M. Fouad BOUKHNIF", "Chef de la Division Gestion"),
                ("M. Noureddine EL FALLAKI", "Chef de la Division Financière et Comptable"),
                ("M. Mohamed ESSALMANI", "Chef de Service Financier"),
                ("Mme Jalila BADRI", "Chef de Service Comptabilité"),
                ("M. Mohamed HAMZAOUI", "Chef de la Division Paiement des Prestations"),
                ("M. Abdelhak JAOUAD", "Chef de Service Centralisation et Suivi"),
                ("M. Brahim NAHI", "Chef de Service Audit"),
                ("Mme Hasnae AIT HAMMOU", "Chef de Service Gouvernance"),
                ("M. Mohamed BESRI", "Cadre au Service Gouvernance")
            ],
            "Cabinet d'audit des comptes": [
                ("M. Khalid FIZAZI", "Managing Partner du Cabinet « FIZAZI »"),
                ("M. Abdelilah ZIAT", "Senior Partner du Cabinet « FIZAZI »")
            ]
        }
        
        # Créer un dictionnaire pour stocker l'état des checkboxes
        if 'participant_checkboxes' not in st.session_state:
            st.session_state.participant_checkboxes = {}
        
        # Afficher les sections et leurs participants
        for section, default_participants in sections.items():
            st.write(f"**{section}**")
            for name, title in default_participants:
                key = f"{name}_{title}"
                if key not in st.session_state.participant_checkboxes:
                    st.session_state.participant_checkboxes[key] = False
                
                if st.checkbox(f"{name} - {title}", key=key, value=st.session_state.participant_checkboxes[key]):
                    st.session_state.participant_checkboxes[key] = True
                    participants.append((name, title, section))
                else:
                    st.session_state.participant_checkboxes[key] = False
        
        # Option pour ajouter des participants supplémentaires
        if st.button("Ajouter un participant"):
            st.session_state.additional_participants.append(len(st.session_state.additional_participants))
        
        # Afficher les champs pour les participants supplémentaires
        for i in st.session_state.additional_participants:
            st.write("**Participant supplémentaire**")
            col1, col2, col3 = st.columns(3)
            with col1:
                name = st.text_input(f"Nom {i+1}")
            with col2:
                title = st.text_input(f"Titre {i+1}")
            with col3:
                section = st.selectbox(f"Section {i+1}", 
                    ["Membres du Comité d'Audit", 
                     "Caisse Marocaine des Retraites",
                     "Cabinet d'audit des comptes"])
            if name and title:
                participants.append((name, title, section))

    # Organiser les participants par section
    participants_by_section = {}
    for name, title, section in participants:
        if section not in participants_by_section:
            participants_by_section[section] = []
        participants_by_section[section].append((name, title))

    # Créer le texte formaté des participants
    participants_text = []
    for section in ["Membres du Comité d'Audit", "Caisse Marocaine des Retraites", "Cabinet d'audit des comptes"]:
        if section in participants_by_section and participants_by_section[section]:
            participants_text.append(f"\n{section}")
            for name, title in participants_by_section[section]:
                participants_text.append(f"- {name}: {title}")

    # Stocker les informations de la réunion
    st.session_state.meeting_info = {
        'pv_number': pv_number,
        'date': date.strftime("%d/%m/%Y"),
        'lieu': lieu,
        'heure_debut': heure_debut.strftime("%H:%M"),
        'heure_fin': heure_fin.strftime("%H:%M"),
        'participants': [(name, title) for name, title, _ in participants],
        'participants_by_section': participants_by_section,
        'participants_text': "\n".join(participants_text)
    }

    # Section d'upload des fichiers
    st.markdown("### 📁 Importation des documents")
    
    # Style CSS pour contrôler individuellement chaque drag and drop
    st.markdown("""
        <style>
        /* Styles de base pour tous les uploaders */
        .stFileUploader > div {
            display: flex;
            align-items: center;
            justify-content: center;
        }
        
        /* Style spécifique pour l'uploader vidéo */
        [data-testid="stFileUploader"]:has(#video_uploader) {
            height: 150px;
            margin-top: 0.5rem;
            margin-bottom: 1rem;
            background-color: rgba(255, 255, 255, 0.05);
        }
        
        /* Style spécifique pour l'uploader d'images */
        [data-testid="stFileUploader"]:has(#image_uploader) {
            height: 180px;
            margin-top: 0.5rem;
            margin-bottom: 1rem;
            background-color: rgba(255, 255, 255, 0.05);
        }
        
        /* Style spécifique pour l'uploader PDF */
        [data-testid="stFileUploader"]:has(#pdf_uploader) {
            height: 160px;
            margin-top: 0.75rem;
            margin-bottom: 1rem;
            background-color: rgba(255, 255, 255, 0.05);
        }
        
        /* Style pour le conteneur des colonnes */
        .row-widget.stHorizontalBlock {
            align-items: flex-start;
            gap: 1.5rem;
        }
        
        /* Style pour les titres des sections */
        .element-container h3 {
            margin-bottom: 0.75rem;
        }
        </style>
    """, unsafe_allow_html=True)
    


    # Créer d'abord les titres dans une rangée
    title_cols = st.columns(3)
    with title_cols[0]:
        st.markdown("### 🎥 Vidéo de la réunion")
    with title_cols[1]:
        st.markdown("### 📝 Images manuscrites")
    with title_cols[2]:
        st.markdown("### 📄 Documents PDF")

    # Ensuite, créer les options radio pour la vidéo dans une rangée séparée
    radio_col, empty_col1, empty_col2 = st.columns(3)
    with radio_col:
        video_upload_mode = st.radio(
            "Mode d'importation :",("Uploader un fichier", "Fournir un lien"),
            horizontal=True,
            key="video_mode"
        )

    # Ensuite, créer les textes d'instructions dans une rangée séparée
    text_cols = st.columns(3)
    with text_cols[0]:
        if video_upload_mode == "Uploader un fichier":
            st.markdown("Importez votre vidéo")
    with text_cols[1]:
        st.markdown("Importez vos images")
    with text_cols[2]:
        st.markdown("Importez vos documents")

    # Enfin, créer les zones de téléchargement dans une rangée séparée
    upload_cols = st.columns(3)
    with upload_cols[0]:
        if video_upload_mode == "Uploader un fichier":
            video_file = st.file_uploader(
                "Importer une vidéo",
                type=["mp4", "vro", "mpeg4"],
                help="Formats acceptés : MP4, VRO, MPEG4 • Limite : 2GB",
                key="video_uploader",
                label_visibility="collapsed"
            )
        else:
            video_url = st.text_input(
                "Lien de la vidéo",
                placeholder="https://drive.google.com/file/d/...",
                help="Lien Google Drive partagé"
            )
    with upload_cols[1]:
        image_files = st.file_uploader(
            "Importer des images",
            type=["jpg", "jpeg", "png"],
            accept_multiple_files=True,
            help="Formats acceptés : JPG, JPEG, PNG • Limite : 2GB par fichier",
            key="image_uploader",
            label_visibility="collapsed"
        )
    with upload_cols[2]:
        pdf_files = st.file_uploader(
            "Importer des PDF",
            type=["pdf"],
            accept_multiple_files=True,
            help="Format accepté : PDF • Limite : 2GB par fichier",
            key="pdf_uploader",
            label_visibility="collapsed"
        )

    # Bouton de démarrage centré avec espace au-dessus
    st.markdown("<div style='text-align: center; margin-top: 2em;'>", unsafe_allow_html=True)
    if st.button("🚀 Démarrer le traitement", use_container_width=True):
        if not st.session_state.meeting_info:
            st.error("❌ Veuillez remplir les informations de base du PV avant de commencer le traitement.")
            return
            
        # Créer des conteneurs pour les résultats
        video_container = st.container()
        images_container = st.container()
        pdfs_container = st.container()
        pv_container = st.container()

        # Traitement de la vidéo
        with video_container:
            st.subheader("🎥 Traitement de la vidéo")
            if (video_file is not None) or (video_url is not None and video_url.strip() != ""):
          # Vérifier si on a soit un fichier soit une URL valide
                with st.spinner("Transcription en cours..."):
                    if video_file:
                        st.session_state.video_transcript = transcribe_video(video_file)
                    elif video_url and video_url.strip():
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".mp4") as temp_video:
                            temp_video_path = temp_video.name
                            if download_video_from_drive(video_url, temp_video_path):
                                if verify_video_file(temp_video_path):
                                    st.session_state.video_transcript = transcribe_video(open(temp_video_path, "rb"))
                                else:
                                    st.error("❌ Le fichier vidéo téléchargé n'est pas valide")
                            else:
                                st.error("❌ Échec du téléchargement de la vidéo")

                    if st.session_state.video_transcript:
                        st.success("✅ Transcription terminée!")
                        st.text_area("Transcription:", st.session_state.video_transcript, height=200)
            else:
                st.info("ℹ️ Aucune vidéo n'a été fournie")

        # Traitement des images
        if image_files:
            with images_container:
                st.subheader("🖼️ Traitement des images")
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                transcriptions = {}
                for idx, image_file in enumerate(image_files):
                    try:
                        status_text.text(f"Analyse de l'image {idx + 1}/{len(image_files)}: {image_file.name}")
                        image_bytes = image_file.read()
                        image = Image.open(io.BytesIO(image_bytes))
                        
                        col1, col2 = st.columns([1, 1])
                        with col1:
                            st.image(image, caption=f"Image {idx + 1}: {image_file.name}", use_column_width=True)
                        with col2:
                            transcription = process_handwritten_image(image_bytes)
                            if transcription:
                                st.text_area(f"Texte reconnu - Image {idx + 1}", transcription, height=150)
                                transcriptions[image_file.name] = transcription
                        
                        progress_bar.progress((idx + 1)/len(image_files))
                    except Exception as e:
                        st.error(f"❌ Erreur lors du traitement de l'image {image_file.name}: {str(e)}")
                
                if transcriptions:
                    st.session_state.handwritten_text = "\n\n".join([f"[Image: {name}]\n{text}" for name, text in transcriptions.items()])
                    st.success("✅ Traitement des images terminé!")

        # Traitement des PDFs
        if pdf_files:
            with pdfs_container:
                st.subheader("📄 Traitement des PDFs")
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                # Réinitialiser les données PDF de la session
                st.session_state.pdf_data = {}
                pdf_summaries_list = [] # Liste temporaire pour l'ancien format
                
                for idx, pdf_file in enumerate(pdf_files):
                    try:
                        status_text.text(f"Analyse du document {idx + 1}/{len(pdf_files)}: {pdf_file.name}")
                        # Lire à nouveau car le pointeur peut être à la fin après l'upload
                        pdf_file.seek(0) 
                        pdf_result = process_pdf(pdf_file)
                        
                        if pdf_result["summary"]:
                            # Stocker le résultat structuré
                            st.session_state.pdf_data[pdf_file.name] = pdf_result
                            # Ajouter au résumé global pour generate_meeting_minutes
                            pdf_summaries_list.append(f"[Document: {pdf_file.name}]\n{pdf_result['summary']}")
                            
                            # Afficher l'aperçu
                            with st.expander(f"📄 Document {idx + 1}: {pdf_file.name} (Analysé)"):
                                st.text_area("Aperçu du contenu extrait:", pdf_result["summary"], height=200)
                                if pdf_result["acronyms"]:
                                    st.write("**Acronymes détectés:**")
                                    st.json(pdf_result["acronyms"])
                                else:
                                    st.write("Aucun acronyme détecté.")
                        else:
                             st.warning(f"Aucun contenu extrait pour {pdf_file.name}")
                        
                        progress_bar.progress((idx + 1)/len(pdf_files))
                    except Exception as e:
                        st.error(f"❌ Erreur lors de l'analyse du PDF {pdf_file.name}: {str(e)}")
                        # Stocker une indication d'erreur
                        st.session_state.pdf_data[pdf_file.name] = {"summary": f"[Erreur: {str(e)}]", "acronyms": {}} 
                        pdf_summaries_list.append(f"[Document: {pdf_file.name}]\n[Erreur lors de l'analyse: {str(e)}]")
                
                # Mettre à jour l'ancien état pdf_summary pour generate_meeting_minutes
                st.session_state.pdf_summary = "\n\n".join(pdf_summaries_list)
                if st.session_state.pdf_data:
                    st.success("✅ Traitement des PDFs terminé!")
                else:
                    st.warning("Aucun PDF n'a pu être traité.")

        # Génération du PV
        with pv_container:
            st.subheader("Génération du PV")
            # Préparer le résumé combiné pour generate_meeting_minutes
            pdf_summary_for_generation = "\n\n".join(
                [f"[Document: {name}]\n{data.get('summary', '')}" 
                 for name, data in st.session_state.get('pdf_data', {}).items()]
            )
            
            if any([st.session_state.video_transcript, st.session_state.handwritten_text, pdf_summary_for_generation]):
                with st.spinner("Génération du PV en cours..."):
                    pv = generate_meeting_minutes(
                        st.session_state.video_transcript,
                        st.session_state.handwritten_text,
                        pdf_summary_for_generation, # Utilise le résumé agrégé
                        st.session_state.meeting_info
                    )
                    if pv:
                        st.success("✅ PV généré avec succès!")
                        st.text_area("Procès-verbal de la réunion:", pv, height=500)
                        
                        # Création et téléchargement du document Word
                        try:
                            doc_buffer = create_word_pv(pv, "logo.png")
                            st.download_button(
                                label="📎 Télécharger le PV en format Word",
                                data=doc_buffer,
                                file_name=f"PV_{st.session_state.meeting_info.get('pv_number', 'NA').replace('/', '_')}_Comite_Audit.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        except Exception as e_word:
                            st.error(f"❌ Erreur lors de la création du document Word: {str(e_word)}")
            else:
                st.warning("⚠️ Aucun contenu à traiter pour générer le PV")

if __name__ == "__main__":
    main() 
